"""Generate a report-ready .docx file for conversation evaluation results."""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.shared import Pt

from .utils import ARTIFACTS_DIR, DATA_DIR, REPORTS_DIR, ensure_dirs, load_json


def _add_table(document: Document, rows: List[Dict[str, Any]], columns: List[str], title: str) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col

    for row in rows:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            value = row.get(col, "")
            cells[i].text = "" if value is None else str(value)


def _qualitative_lines(per_sample: List[Dict[str, Any]]) -> List[str]:
    lines: List[str] = []
    by_model: Dict[str, List[Dict[str, Any]]] = {}
    for r in per_sample:
        key = f"{r['provider']}::{r['model']}"
        by_model.setdefault(key, []).append(r)

    for key, rows in by_model.items():
        good = [r for r in rows if not r.get("error")]
        if not good:
            lines.append(f"{key}: no successful generations due to provider/runtime errors.")
            continue

        faithful = sum(r["faithful"] for r in good) / len(good)
        avg_rouge = sum(float(r["rougeL"]) for r in good) / len(good)
        adv = [r for r in good if not r["expected_in_corpus"]]
        adv_abstain = (sum(1 for r in adv if r["abstained"]) / len(adv)) if adv else 0.0

        if faithful >= 0.7:
            lines.append(f"{key} stays mostly grounded in retrieved evidence, with strong faithfulness behavior.")
        else:
            lines.append(f"{key} shows frequent grounding issues; unsupported claims appear in several answers.")

        if avg_rouge >= 0.2:
            lines.append(f"{key} tracks reference content reasonably well on lexical overlap (ROUGE-L).")
        else:
            lines.append(f"{key} often diverges from reference phrasing and may under-cover key points.")

        lines.append(f"{key} abstains on adversarial questions in {adv_abstain * 100:.1f}% of cases.")

    return lines


def generate_report(dataset_path: str, raw_path: str, aggregate_path: str, per_sample_path: str, output_name: str) -> Path:
    ensure_dirs()
    dataset = load_json(DATA_DIR / dataset_path)
    raw = load_json(ARTIFACTS_DIR / raw_path)
    aggregate = load_json(ARTIFACTS_DIR / aggregate_path)["rows"]
    per_sample = load_json(ARTIFACTS_DIR / per_sample_path)["rows"]

    document = Document()
    base_style = document.styles["Normal"]
    base_style.font.name = "Calibri"
    base_style.font.size = Pt(11)

    document.add_heading("Conversation System Evaluation Report", level=1)
    document.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph("Corpus: r/cscareerquestions extracted Reddit repository (career.db + RAG index).")

    document.add_heading("Assumptions", level=2)
    assumptions = [
        "Reference answers for in-corpus items are deterministic extractive summaries from retrieved chunks, not human-written abstractive gold answers.",
        "Faithfulness is computed using a deterministic heuristic: lexical support ratio + unsupported sentence ratio against retrieved context.",
        "Adversarial questions are considered correctly handled when the model abstains or explicitly states missing evidence.",
    ]
    for line in assumptions:
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("Provider Run Status", level=2)
    for provider, info in raw.get("provider_summary", {}).items():
        document.add_paragraph(
            f"{provider}: available={info.get('available')} success={info.get('num_success')} failed={info.get('num_failed')} missing_key={info.get('missing_key')}"
        )

    _add_table(
        document,
        aggregate,
        [
            "provider",
            "model",
            "num_questions",
            "num_success",
            "num_failed",
            "bleu",
            "rougeL",
            "bertscore_f1",
            "faithfulness_pct",
            "avg_support_ratio",
            "adversarial_refusal_pct",
        ],
        "Model Comparison",
    )

    trimmed_rows = []
    for r in per_sample:
        trimmed_rows.append(
            {
                "question_id": r["question_id"],
                "provider": r["provider"],
                "model": r["model"],
                "question_type": r["question_type"],
                "expected_in_corpus": r["expected_in_corpus"],
                "bleu": r["bleu"],
                "rougeL": r["rougeL"],
                "bertscore_f1": r["bertscore_f1"],
                "faithful": r["faithful"],
                "abstained": r["abstained"],
                "error": r.get("error") or "",
            }
        )

    _add_table(
        document,
        trimmed_rows,
        [
            "question_id",
            "provider",
            "model",
            "question_type",
            "expected_in_corpus",
            "bleu",
            "rougeL",
            "bertscore_f1",
            "faithful",
            "abstained",
            "error",
        ],
        "Per-Question Breakdown",
    )

    document.add_heading("Qualitative Analysis", level=2)
    for line in _qualitative_lines(per_sample):
        document.add_paragraph(line, style="List Bullet")

    document.add_heading("Adversarial Question Analysis", level=2)
    adv_ids = [q["id"] for q in dataset["questions"] if not q["expected_in_corpus"]]
    adv_rows = [r for r in per_sample if r["question_id"] in adv_ids]
    for r in adv_rows:
        document.add_paragraph(
            f"{r['question_id']} | {r['provider']}::{r['model']} | abstained={r['abstained']} | faithful={r['faithful']} | error={r.get('error') or 'none'}"
        )

    document.add_heading("Key Observations", level=2)
    document.add_paragraph("Faithfulness percentages are interpreted jointly with support-ratio diagnostics to detect unsupported claims.", style="List Bullet")
    document.add_paragraph("Adversarial behavior is explicitly measured through abstention/refusal signal, not only lexical metrics.", style="List Bullet")
    document.add_paragraph("BLEU/ROUGE are useful but limited for Reddit-style open-ended answers, so BERTScore and faithfulness are primary indicators.", style="List Bullet")

    out_path = REPORTS_DIR / output_name
    document.save(out_path)
    return out_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate conversation evaluation DOCX report")
    parser.add_argument("--dataset", default="conversation_eval_questions.json")
    parser.add_argument("--raw", default="raw_model_outputs.json")
    parser.add_argument("--aggregate", default="aggregate_metrics.json")
    parser.add_argument("--per-sample", default="per_sample_scores.json")
    parser.add_argument("--output", default="conversation_system_evaluation_report.docx")
    args = parser.parse_args()

    out = generate_report(args.dataset, args.raw, args.aggregate, args.per_sample, args.output)
    print(f"Report generated: {out}")


if __name__ == "__main__":
    main()
